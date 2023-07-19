from docxtpl import DocxTemplate
from sympy import symbols, N, diff
import sympy
import streamlit as st
import matplotlib.pyplot as plt
from docx.shared import Inches


def bi_section(f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond):
    x, e, pi = symbols('x e pi')
    f_x = N(f_x.subs(e, sympy.E))
    f_x = N(f_x.subs(pi, sympy.pi))

    f_x_l = N(f_x.subs(x, x_lower))
    f_x_u = N(f_x.subs(x, x_upper))

    x_root = (x_upper + x_lower) / 2
    f_x_r = N(f_x.subs(x, x_root))

    data = {
        'x lower': [],
        'x upper': [],
        'x root': [],
        'f(xl)': [],
        'f(xr)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    if true_relative_error_cond is not None:
        true_relative_error = abs((x_exact - x_root) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while (true_relative_error_cond is None or true_relative_error > true_relative_error_cond) and (
            approximate_relative_error_cond is None or approximate_relative_error > approximate_relative_error_cond) and (
            iter_num_cond is None or iter_num <= iter_num_cond):
        data['x lower'].append(x_lower)
        data['x upper'].append(x_upper)
        data['x root'].append(x_root)
        data['f(xl)'].append(f_x_l)
        data['f(xr)'].append(f_x_r)
        data['Et(%)'].append(true_relative_error if true_relative_error_cond is not None else '---')
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        if f_x_r * f_x_u < 0:
            x_lower = x_root
            f_x_l = f_x_r
        else:
            x_upper = x_root
            f_x_u = f_x_r

        pre_x_r = x_root
        x_root = float((x_upper + x_lower) / 2)
        f_x_r = N(f_x.subs(x, x_root))

        iter_num += 1
        if true_relative_error_cond is not None:
            true_relative_error = abs((x_exact - x_root) / x_exact) * 100
        approximate_relative_error = abs((x_root - pre_x_r) / x_root) * 100

        errors[iter_num] = [true_relative_error if true_relative_error_cond is not None else None,
                            approximate_relative_error]

    return data, errors


def false_position(f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond):
    x, e, pi = symbols('x e pi')
    f_x = N(f_x.subs(e, sympy.E))
    f_x = N(f_x.subs(pi, sympy.pi))

    f_x_l = N(f_x.subs(x, x_lower))
    f_x_u = N(f_x.subs(x, x_upper))

    x_root = x_upper - f_x_u * (x_upper - x_lower) / (f_x_u - f_x_l)
    f_x_r = N(f_x.subs(x, x_root))

    data = {
        'x lower': [],
        'x upper': [],
        'x root': [],
        'f(xl)': [],
        'f(xu)': [],
        'f(xr)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    if true_relative_error_cond is not None:
        true_relative_error = abs((x_exact - x_root) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while (true_relative_error_cond is None or true_relative_error > true_relative_error_cond) and (
            approximate_relative_error_cond is None or approximate_relative_error > approximate_relative_error_cond) and (
            iter_num_cond is None or iter_num <= iter_num_cond):
        data['x lower'].append(x_lower)
        data['x upper'].append(x_upper)
        data['x root'].append(x_root)
        data['f(xr)'].append(f_x_r)
        data['f(xl)'].append(f_x_l)
        data['f(xu)'].append(f_x_u)
        data['Et(%)'].append(true_relative_error if true_relative_error_cond is not None else '---')
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        if f_x_r * f_x_u < 0:
            x_lower = x_root
            f_x_l = f_x_r
        else:
            x_upper = x_root
            f_x_u = f_x_r

        pre_x_r = x_root
        x_root = float(x_upper - f_x_u * (x_upper - x_lower) / (f_x_u - f_x_l))
        f_x_r = N(f_x.subs(x, x_root))

        iter_num += 1
        if true_relative_error_cond is not None:
            true_relative_error = abs((x_exact - x_root) / x_exact) * 100
        approximate_relative_error = abs((x_root - pre_x_r) / x_root) * 100

        errors[iter_num] = [true_relative_error if true_relative_error_cond is not None else None,
                            approximate_relative_error]

    return data, errors


def fixed_point(f_x, g_x, x_exact, x_i, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond):
    x, e, pi = symbols('x e pi')
    f_x = N(f_x.subs(e, sympy.E))
    f_x = N(f_x.subs(pi, sympy.pi))

    g_x = N(g_x.subs(e, sympy.E))
    g_x = N(g_x.subs(pi, sympy.pi))

    f_x_i = N(f_x.subs(x, x_i))
    g_x_i = N(g_x.subs(x, x_i))

    data = {
        'x': [],
        'f(x)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    if true_relative_error_cond is not None:
        true_relative_error = abs((x_exact - x_i) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while (true_relative_error_cond is None or true_relative_error > true_relative_error_cond) and (
            approximate_relative_error_cond is None or approximate_relative_error > approximate_relative_error_cond) and (
            iter_num_cond is None or iter_num <= iter_num_cond):
        data['x'].append(x_i)
        data['f(x)'].append(f_x_i)
        data['Et(%)'].append(true_relative_error if true_relative_error_cond is not None else '---')
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        pre_x_i = x_i

        x_i = g_x_i
        g_x_i = float(N(g_x.subs(x, x_i)))
        f_x_i = float(N(f_x.subs(x, x_i)))

        iter_num += 1
        if true_relative_error_cond is not None:
            true_relative_error = abs((x_exact - x_i) / x_exact) * 100
        approximate_relative_error = abs((x_i - pre_x_i) / x_i) * 100

        errors[iter_num] = [true_relative_error if true_relative_error_cond is not None else None,
                            approximate_relative_error]

    return data, errors


def newton_raphson(f_x, x_exact, x_i, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond):
    x, e, pi = symbols('x e pi')
    f_x = N(f_x.subs(e, sympy.E))
    f_x = N(f_x.subs(pi, sympy.pi))

    f_x_i = f_x.subs(x, x_i)
    x_i_1 = x_i - f_x.subs(x, x_i) / diff(f_x, x, 1).subs(x, x_i)

    data = {
        'x': [],
        'f(x)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    if true_relative_error_cond is not None:
        true_relative_error = abs((x_exact - x_i) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while (true_relative_error_cond is None or true_relative_error > true_relative_error_cond) and (approximate_relative_error_cond is None or approximate_relative_error > approximate_relative_error_cond) and (iter_num_cond is None or iter_num <= iter_num_cond):
        data['x'].append(x_i)
        data['f(x)'].append(f_x_i)
        data['Et(%)'].append(true_relative_error if true_relative_error_cond is not None else '---')
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        pre_x_i = x_i

        x_i = x_i_1
        f_x_i = f_x.subs(x, x_i)
        x_i_1 = float(x_i - f_x.subs(x, x_i) / diff(f_x, x, 1).subs(x, x_i))

        iter_num += 1
        if true_relative_error_cond is not None:
            true_relative_error = abs((x_exact - x_i) / x_exact) * 100
        approximate_relative_error = abs((x_i - pre_x_i) / x_i) * 100

        errors[iter_num] = [true_relative_error if true_relative_error_cond is not None else None, approximate_relative_error]

    return data, errors


def secant(f_x, x_exact, x_i, x_i_1, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond):
    x, e, pi = symbols('x e pi')
    f_x = N(f_x.subs(e, sympy.E))
    f_x = N(f_x.subs(pi, sympy.pi))

    f_x_i = f_x.subs(x, x_i)
    f_x_i_1 = f_x.subs(x, x_i_1)

    x_i_2 = x_i_1 - (f_x.subs(x, x_i_1) * (x_i - x_i_1)) / (f_x.subs(x, x_i) - f_x.subs(x, x_i_1))

    data = {
        'Xi-1': [],
        'Xi': [],
        'Xi+1': [],
        'f(Xi-1)': [],
        'f(Xi)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    if true_relative_error_cond is not None:
        true_relative_error = abs((x_exact - x_i_2) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while (true_relative_error_cond is None or true_relative_error > true_relative_error_cond) and (approximate_relative_error_cond is None or approximate_relative_error > approximate_relative_error_cond) and (iter_num_cond is None or iter_num <= iter_num_cond):
        data['Xi-1'].append(x_i)
        data['Xi'].append(x_i_1)
        data['Xi+1'].append(x_i_2)
        data['f(Xi-1)'].append(f_x_i)
        data['f(Xi)'].append(f_x_i_1)
        data['Et(%)'].append(true_relative_error if true_relative_error_cond is not None else '---')
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        x_i = x_i_1
        x_i_1 = x_i_2
        x_i_2 = float(x_i_1 - (f_x.subs(x, x_i_1) * (x_i - x_i_1)) / (f_x.subs(x, x_i) - f_x.subs(x, x_i_1)))

        f_x_i = f_x.subs(x, x_i)
        f_x_i_1 = f_x.subs(x, x_i_1)

        iter_num += 1
        if true_relative_error_cond is not None:
            true_relative_error = abs((x_exact - x_i_2) / x_exact) * 100
        approximate_relative_error = abs((x_i_2 - x_i_1) / x_i_2) * 100

        errors[iter_num] = [true_relative_error if true_relative_error_cond is not None else None, approximate_relative_error]

    return data, errors


def draw_approximate_and_true_errors_for_single_method(method_name, errors):
    plt.plot(errors.keys(), [error[0] for error in errors.values()], label="Et(%)")
    plt.plot(errors.keys(), [error[1] for error in errors.values()], label="Ea(%)")
    plt.legend()

    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')

    plt.title(f'True and approximate errors for {method_name} method')

    plt.savefig(f'{method_name}.png')
    st.image(f'{method_name}.png', use_column_width=True)

    plt.close()


def draw_true_errors_for_all_methods(bi_section_errors, false_position_errors, fixed_point_errors, newton_raphson_errors, secant_errors):
    if bi_section_errors is not None:
        plt.plot(bi_section_errors.keys(), [error[0] for error in bi_section_errors.values()], label="Et(%) for BiSection")
    if false_position_errors is not None:
        plt.plot(false_position_errors.keys(), [error[0] for error in false_position_errors.values()], label="Et(%) for False Position")
    if fixed_point_errors is not None:
        plt.plot(fixed_point_errors.keys(), [error[0] for error in fixed_point_errors.values()], label="Et(%) for Fixed Point")
    if newton_raphson_errors is not None:
        plt.plot(newton_raphson_errors.keys(), [error[0] for error in newton_raphson_errors.values()], label="Et(%) for Newton Raphson")
    if secant_errors is not None:
        plt.plot(secant_errors.keys(), [error[0] for error in secant_errors.values()], label="Et(%) for Secant")
    plt.legend()
    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')
    plt.title('True errors')

    plt.savefig('chart.png')
    st.pyplot(plt)
    plt.close()


def draw_approximate_errors_for_all_methods(bi_section_errors, false_position_errors, fixed_point_errors, newton_raphson_errors, secant_errors):
    if bi_section_errors is not None:
        plt.plot(bi_section_errors.keys(), [error[1] for error in bi_section_errors.values()], label="Ea(%) for BiSection")
    if false_position_errors is not None:
        plt.plot(false_position_errors.keys(), [error[1] for error in false_position_errors.values()], label="Ea(%) for False Position")
    if fixed_point_errors is not None:
        plt.plot(fixed_point_errors.keys(), [error[1] for error in fixed_point_errors.values()], label="Ea(%) for Fixed Point")
    if newton_raphson_errors is not None:
        plt.plot(newton_raphson_errors.keys(), [error[1] for error in newton_raphson_errors.values()], label="Ea(%) for Newton Raphson")
    if secant_errors is not None:
        plt.plot(secant_errors.keys(), [error[1] for error in secant_errors.values()], label="Ea(%) for Secant")
    plt.legend()
    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')
    plt.title('Approximate errors')

    plt.savefig('chart.png')
    st.pyplot(plt)
    plt.close()


def get_bisection_doc(f_x, iteration, approximate_cond, xexact, data):
    main_path = r"templetes/bisection.docx"
    template = DocxTemplate(main_path)

    context = {
        "function": f_x,
        "iteration": iteration,
        "approximate": approximate_cond,
        "xt": xexact,

        "method":'BiSection',
        "chapter":'Chapter 2',

        "xl1": round(data['x lower'][0], 5),
        "xu1": round(data['x upper'][0], 5),
        "xr1": round(data['x root'][0], 5),
        "fxl1": round(data['f(xl)'][0], 5),
        "fxr1": round(data['f(xr)'][0], 5),
        "multi1": round(data['f(xl)'][0] * data['f(xr)'][0], 5),
        "operator1": ">" if (data['f(xl)'][0] * data['f(xr)'][0] > 0) else "<",
        "newX1": "xl" if (data['f(xl)'][0] * data['f(xr)'][0] > 0) else "xu",
        "ea1": data['Ea(%)'][0],

        "xl2": round(data['x lower'][1], 5),
        "xu2": round(data['x upper'][1], 5),
        "xr2": round(data['x root'][1], 5),
        "fxl2": round(data['f(xl)'][1], 5),
        "fxr2": round(data['f(xr)'][1], 5),
        "multi2": round(data['f(xl)'][1] * data['f(xr)'][1], 5),
        "operator2": ">" if (data['f(xl)'][1] * data['f(xr)'][1] > 0) else "<",
        "newX2": "xl" if (data['f(xl)'][1] * data['f(xr)'][1] > 0) else "xu",
        "ea2": round(data['Ea(%)'][1], 5),

        "xl3": round(data['x lower'][2], 5),
        "xu3": round(data['x upper'][2], 5),
        "xr3": round(data['x root'][2], 5),
        "fxl3": round(data['f(xl)'][2], 5),
        "fxr3": round(data['f(xr)'][2], 5),
        "multi3": round(data['f(xl)'][2] * data['f(xr)'][2], 5),
        "operator3": ">" if (data['f(xl)'][2] * data['f(xr)'][2] > 0) else "<",
        "newX3": "xl" if (data['f(xl)'][2] * data['f(xr)'][2] > 0) else "xu",
        "ea3": round(data['Ea(%)'][2], 5),

        "x_last": round(data['x root'][-1], 5),
    }

    template.render(context)

    headers = ['iteration'] + list(data.keys())
    table = template.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    template.add_picture('Bi Section.png', width=Inches(5))

    template.save('BiSection.docx')


def get_false_position_doc(f_x, iteration, approximate_cond, xexact, data):
    main_path = r"templetes/false_position.docx"
    template = DocxTemplate(main_path)

    context = {
        "function": f_x,
        "iteration": iteration,
        "approximate": approximate_cond,
        "xt": xexact,

        "method":'False Position',
        "chapter":'Chapter 2',

        "xl1": round(data['x lower'][0], 5),
        "xu1": round(data['x upper'][0], 5),
        "xr1": round(data['x root'][0], 5),
        "fxl1": round(data['f(xl)'][0], 5),
        "fxu1": round(data['f(xu)'][0], 5),
        "fxr1": round(data['f(xr)'][0], 5),
        "multi1": round(data['f(xl)'][0] * data['f(xr)'][0], 5),
        "operator1": ">" if (data['f(xl)'][0] * data['f(xr)'][0] > 0) else "<",
        "newX1": "xl" if (data['f(xl)'][0] * data['f(xr)'][0] > 0) else "xu",
        "ea1": data['Ea(%)'][0],

        "xl2": round(data['x lower'][1], 5),
        "xu2": round(data['x upper'][1], 5),
        "xr2": round(data['x root'][1], 5),
        "fxl2": round(data['f(xl)'][1], 5),
        "fxu2": round(data['f(xu)'][1], 5),
        "fxr2": round(data['f(xr)'][1], 5),
        "multi2": round(data['f(xl)'][1] * data['f(xr)'][1], 5),
        "operator2": ">" if (data['f(xl)'][1] * data['f(xr)'][1] > 0) else "<",
        "newX2": "xl" if (data['f(xl)'][1] * data['f(xr)'][1] > 0) else "xu",
        "ea2": round(data['Ea(%)'][1], 5),

        "xl3": round(data['x lower'][2], 5),
        "xu3": round(data['x upper'][2], 5),
        "xr3": round(data['x root'][2], 5),
        "fxl3": round(data['f(xl)'][2], 5),
        "fxu3": round(data['f(xu)'][2], 5),
        "fxr3": round(data['f(xr)'][2], 5),
        "multi3": round(data['f(xl)'][2] * data['f(xr)'][2], 5),
        "operator3": ">" if (data['f(xl)'][2] * data['f(xr)'][2] > 0) else "<",
        "newX3": "xl" if (data['f(xl)'][2] * data['f(xr)'][2] > 0) else "xu",
        "ea3": round(data['Ea(%)'][2], 5),

        "x_last": round(data['x root'][-1], 5),
    }

    template.render(context)

    headers = ['iteration'] + list(data.keys())
    table = template.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    template.add_picture('False Position.png', width=Inches(5))

    template.save('False_Position.docx')


def get_fixed_point_doc(f_x, iteration, approximate_cond, data):
    main_path = r"templetes/fixed_point.docx"
    template = DocxTemplate(main_path)

    context = {
        "function": f_x,
        "iteration": iteration,
        "approximate": approximate_cond,

        "method":'Fixed Point',
        "chapter":'Chapter 2',

        "x": round(data['x'][0], 5),
        "gx": round(data['f(x)'][0], 5),
        "x1": round(data['x'][1], 5),
        "ea1": data['Ea(%)'][0],
        "x2": round(data['x'][2], 5),
        "ea2": round(data['Ea(%)'][1], 5),

        "x3": round(data['x'][3], 5),
        "ea3": round(data['Ea(%)'][2], 5),
        "xreal": round(data['x'][-1], 5),
    }

    template.render(context)

    headers = ['iteration'] + list(data.keys())
    table = template.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    template.add_picture('Fixed Point.png', width=Inches(5))

    template.save('FixedPoint.docx')


def get_newton_raphson_doc(f_x, iteration, approximate_cond, data):
    main_path = r"templetes/newton_raphson.docx"
    template = DocxTemplate(main_path)

    x = symbols('x')

    context = {
        "function": f_x,
        "iteration": iteration,
        "approximate": approximate_cond,

        "method": 'Newton Raphson',
        "chapter": 'Chapter 2',

        "x": round(data['x'][0], 5),
        "funct": diff(f_x, x, 1),
        "x1": round(data['x'][1], 5),
        "ea1": data['Ea(%)'][0],
        "x2": round(data['x'][2], 5),
        "ea2": round(data['Ea(%)'][1], 5),
        "x3": round(data['x'][3], 5) if len(data['x']) > 3 else '---',
        "ea3": round(data['Ea(%)'][2], 5) if len(data['x']) > 3 else '---',
        "xreal": round(data['x'][-1], 5),
    }

    template.render(context)

    headers = ['iteration'] + list(data.keys())
    table = template.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    template.add_picture('Newton Raphson.png', width=Inches(5))

    template.save('Newton Raphson.docx')


def get_secant_doc(f_x, iteration, approximate_cond, data):
    main_path = r"templetes/secant.docx"
    template = DocxTemplate(main_path)

    context = {
        "function": f_x,
        "iteration": iteration,
        "approximate": approximate_cond,

        "method": 'Secant',
        "chapter": 'Chapter 2',

        "x": round(data['Xi-1'][0], 5),
        "x1": round(data['Xi'][0], 5),
        "fx": round(data['f(Xi-1)'][0], 5),
        "fx1": round(data['f(Xi)'][0], 5),
        "x2": round(data['Xi+1'][0], 5),
        "ea1": data['Ea(%)'][0],

        "fx2": round(data['f(Xi)'][1], 5),
        "x3": round(data['Xi+1'][1], 5),
        "ea2": round(data['Ea(%)'][1], 5),


        "fx3": round(data['f(Xi)'][2], 5),
        "x4": round(data['Xi+1'][2], 5),
        "ea3": round(data['Ea(%)'][2], 5),
        "xreal": round(data['Xi+1'][-1], 5),
    }

    template.render(context)

    headers = ['iteration'] + list(data.keys())
    table = template.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    template.add_picture('Secant.png', width=Inches(5))

    template.save('Secant.docx')
