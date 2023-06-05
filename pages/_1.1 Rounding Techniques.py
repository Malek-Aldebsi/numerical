import sympy
from sympy import symbols, N, E, sympify
import streamlit as st

from docx import Document

def get_doc(title, f_x, steps, result, Sf, rounding_technique):
    document = Document()

    document.add_heading(title, 0)
    paragraph = document.add_paragraph(style='Intense Quote')
    paragraph.add_run('The solution of this equation:\n').bold = False
    paragraph.add_run(f'"{f_x}"\n').center = True
    paragraph.add_run(f'using ').bold = False
    paragraph.add_run(f'{rounding_technique} ').bold = True
    paragraph.add_run(f'with ').bold = False
    paragraph.add_run(f'{Sf} ').bold = True
    paragraph.add_run(f'significant figures is:\n\n').bold = False
    for index, step in enumerate(steps):
        paragraph.add_run(f'step {index + 1}:\t{step}\n').bold = False
    paragraph.add_run(f'\nfinal result:\t{"".join(result)}').bold = True

    document.save('documentation.docx')


def priorities(f_x, x_1=.0, x_2=.0, x_3=.0, x_4=.0, x_5=.0, x_6=.0, x_7=.0, x_8=.0, x_9=.0, x_10=.0, with_rounding=False, with_chopping=False, Sf=3, is_main=False):
    steps = []
    def sub(expression):
        return N(expression.subs([(x1, x_1), (x2, x_2), (x3, x_3), (x4, x_4), (x5, x_5), (x6, x_6), (x7, x_7), (x8, x_8), (x9, x_9), (x10, x_10), (e, E), (pi, sympy.pi)]))

    def chopping(number, sf):
        number = str(number)
        number = list(number)
        nz_flag = False
        counter = 0
        for i, ch in enumerate(number):
            if ch not in ['0', '.']:
                nz_flag = True
            if nz_flag and ch != '.':
                counter += 1
            if counter > sf and ch != '.':
                number[i] = '0'
        number = ''.join(number)
        number = float(number)
        return number

    def to_correct_sg(num):
        if with_rounding:
            return format(num, f'.{Sf}g')
        elif with_chopping:
            return chopping(num, Sf)
        else:
            return num
    x1, x2, x3, x4, x5, x6, x7, x8, x9, x10, e, pi = symbols('x1 x2 x3 x4 x5 x6 x7 x8 x9 x10 e pi')

    if is_main:
        steps.append(''.join(f_x))

    for element in f_x:
        if element in ['sin', 'cos', 'tan', 'ln']:
            function_index = f_x.index(element)
            open_bracket = f_x[function_index:].index('(') + function_index
            close_bracket = f_x[function_index:].index(')') + function_index
            function = f_x[function_index] + '(' + str(priorities(f_x[open_bracket+1: close_bracket], x_1, x_2, x_3, x_4, x_5, x_6, x_7, x_8, x_9, x_10, with_rounding, with_chopping, Sf)[0]) + ')'
            function = sympify(function)
            function = to_correct_sg(sub(function))
            f_x = f_x[:function_index] + [function] + f_x[close_bracket+1:]
            if is_main:
                steps.append(''.join(f_x))

    while '(' in f_x:
        open_bracket = f_x.index('(')
        close_bracket = f_x.index(')')
        f_x = f_x[:open_bracket] + priorities(f_x[open_bracket+1: close_bracket], x_1, x_2, x_3, x_4, x_5, x_6, x_7, x_8, x_9, x_10, with_rounding, with_chopping, Sf) + f_x[close_bracket + 1:]
        if is_main:
            steps.append(''.join(f_x))

    while '**' in f_x:
        power_index = f_x.index('**')

        first_ele_index = power_index - 1
        first_ele = sympify(f_x[first_ele_index])

        second_ele_index = power_index + 1
        second_ele = sympify(f_x[second_ele_index])

        first_ele = sub(first_ele)
        second_ele = sub(second_ele)

        f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele ** second_ele)] + f_x[second_ele_index + 1:]
        if is_main:
            steps.append(''.join(f_x))

    while '*' in f_x or '/' in f_x:
        if '*' in f_x and '/' in f_x:
            multiply_index = f_x.index('*')
            division_index = f_x.index('/')

            if multiply_index < division_index:
                first_ele_index = multiply_index - 1
                first_ele = sympify(f_x[first_ele_index])

                second_ele_index = multiply_index + 1
                second_ele = sympify(f_x[second_ele_index])

                first_ele = sub(first_ele)
                second_ele = sub(second_ele)

                f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele * second_ele)] + f_x[second_ele_index + 1:]

            elif division_index < multiply_index:
                first_ele_index = division_index - 1
                first_ele = sympify(f_x[first_ele_index])

                second_ele_index = division_index + 1
                second_ele = sympify(f_x[second_ele_index])

                first_ele = sub(first_ele)
                second_ele = sub(second_ele)

                f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele / second_ele)] + f_x[second_ele_index + 1:]

        elif '*' in f_x:
            multiply_index = f_x.index('*')

            first_ele_index = multiply_index - 1
            first_ele = sympify(f_x[first_ele_index])

            second_ele_index = multiply_index + 1
            second_ele = sympify(f_x[second_ele_index])

            first_ele = sub(first_ele)
            second_ele = sub(second_ele)

            f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele * second_ele)] + f_x[second_ele_index + 1:]

        elif '/' in f_x:
            division_index = f_x.index('/')

            first_ele_index = division_index - 1
            first_ele = sympify(f_x[first_ele_index])

            second_ele_index = division_index + 1
            second_ele = sympify(f_x[second_ele_index])

            first_ele = sub(first_ele)
            second_ele = sub(second_ele)


            f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele / second_ele)] + f_x[second_ele_index + 1:]

        if is_main:
            steps.append(''.join(f_x))

    while '+' in f_x or '-' in f_x:
        if '+' in f_x and '-' in f_x:
            summation_index = f_x.index('+')
            subtraction_index = f_x.index('-')

            if summation_index < subtraction_index:
                first_ele_index = summation_index - 1
                first_ele = sympify(f_x[first_ele_index])

                second_ele_index = summation_index + 1
                second_ele = sympify(f_x[second_ele_index])

                first_ele = sub(first_ele)
                second_ele = sub(second_ele)

                f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele + second_ele)] + f_x[second_ele_index + 1:]

            elif subtraction_index < summation_index:
                first_ele_index = subtraction_index - 1
                first_ele = sympify(f_x[first_ele_index])

                second_ele_index = subtraction_index + 1
                second_ele = sympify(f_x[second_ele_index])

                first_ele = sub(first_ele)
                second_ele = sub(second_ele)

                f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele - second_ele)] + f_x[second_ele_index + 1:]

        elif '+' in f_x:
            summation_index = f_x.index('+')

            first_ele_index = summation_index - 1
            first_ele = sympify(f_x[first_ele_index])

            second_ele_index = summation_index + 1
            second_ele = sympify(f_x[second_ele_index])

            first_ele = sub(first_ele)
            second_ele = sub(second_ele)

            f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele + second_ele)] + f_x[second_ele_index + 1:]

        elif '-' in f_x:
            subtraction_index = f_x.index('-')

            first_ele_index = subtraction_index - 1
            first_ele = sympify(f_x[first_ele_index])

            second_ele_index = subtraction_index + 1
            second_ele = sympify(f_x[second_ele_index])

            first_ele = sub(first_ele)
            second_ele = sub(second_ele)

            f_x = f_x[:first_ele_index] + [to_correct_sg(first_ele - second_ele)] + f_x[second_ele_index + 1:]
        if is_main:
            steps.append(''.join(f_x))

    while any(elem in f_x for elem in ['x1', 'x2', 'x3', 'x4', 'x5', 'x6', 'x7', 'x8', 'x9', 'x10', 'e', 'pi']):
        f_x[0] = sub(sympify(f_x[0]))
        if is_main:
            steps.append(''.join(f_x))

    if is_main:
        return f_x, steps
    else:
        return f_x


st.title('Rounding Techniques')

with st.form("Rounding_Techniques_form"):
    f_x = st.text_input("Enter your function", value=r"x1 / x2 + cos ( x3 ) / x4")
    st.latex(f_x)

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        x_1 = st.text_input("value of x1", value=14)
        x_1 = float(x_1)
        st.write('x1:', x_1)

        x_6 = st.text_input("value of x6", value=0)
        x_6 = float(x_6)
        st.write('x6:', x_6)

    with col2:
        x_2 = st.text_input("value of x2", value=13)
        x_2 = float(x_2)
        st.write('x2:', x_2)

        x_7 = st.text_input("value of x7", value=0)
        x_7 = float(x_7)
        st.write('x7:', x_7)

    with col3:
        x_3 = st.text_input("value of x3", value=5)
        x_3 = float(x_3)
        st.write('x3:', x_3)

        x_8 = st.text_input("value of x8", value=0)
        x_8 = float(x_8)
        st.write('x8:', x_8)

    with col4:
        x_4 = st.text_input("value of x4", value=5)
        x_4 = float(x_4)
        st.write('x4:', x_4)

        x_9 = st.text_input("value of x9", value=0)
        x_9 = float(x_9)
        st.write('x9:', x_9)

    with col5:
        x_5 = st.text_input("value of x5", value=0)
        x_5 = float(x_5)
        st.write('x5:', x_5)

        x_10 = st.text_input("value of x10", value=0)
        x_10 = float(x_10)
        st.write('x10:', x_10)

    col1, col2 = st.columns(2)

    with col1:
        options = ['Rounding', 'Chopping']
        selected_option = st.selectbox('Select an option:', options)

    with col2:
        Sf = st.text_input("# of significant figures", value=3)
        Sf = int(Sf)
        st.write('Sf:', Sf)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        f_x = f_x.split()
        result, steps = priorities(f_x, x_1, x_2, x_3, x_4, x_5, x_6, x_7, x_8, x_9, x_10, selected_option=='Rounding', selected_option=='Chopping', Sf, is_main=True)
        for index, step in enumerate(steps):
            st.latex(f'step{index+1}: {step}')
        st.latex(f'final result: {"".join(result)}')

        get_doc('Rounding Techniques', ''.join(f_x), steps, result, Sf, selected_option)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")
