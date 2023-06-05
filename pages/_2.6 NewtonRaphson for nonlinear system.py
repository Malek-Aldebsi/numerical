from sympy import symbols, sympify, diff, N, E, Matrix

import streamlit as st
import pandas as pd

import matplotlib.pyplot as plt
import matplotlib

from docx import Document
from docx.shared import Inches


def draw_approximate_and_true_errors(errors):
    plt.plot(errors.keys(), [error[0] for error in errors.values()], label="Et(%)(x)")
    plt.plot(errors.keys(), [error[1] for error in errors.values()], label="Et(%)(y)")
    plt.plot(errors.keys(), [error[2] for error in errors.values()], label="Ea(%)(x)")
    plt.plot(errors.keys(), [error[3] for error in errors.values()], label="Ea(%)(y)")
    plt.legend()

    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')

    plt.title('True and approximate errors for Newton Raphson non linear system method')

    plt.savefig('chart.png')
    st.pyplot(plt)


def newton_raphson_for_non_linear_system(f_1, f_2, x_exact, y_exact, x_i, y_i, approximate_relative_error_cond_x=0.005, approximate_relative_error_cond_y=0.005, true_relative_error_cond_x=0.005, true_relative_error_cond_y=0.005, iter_num_cond=10):
    x, y, e = symbols('x y e')
    f_1 = N(f_1.subs(e, E))
    f_2 = N(f_2.subs(e, E))

    X = Matrix([[x_i], [y_i]])
    F = Matrix([[N(f_1.subs({x: X[0], y: X[1]}))], [N(f_2.subs({x: X[0], y: X[1]}))]])

    df1_dx = diff(f_1, x, 1)
    df1_dy = diff(f_1, y, 1)
    df2_dx = diff(f_2, x, 1)
    df2_dy = diff(f_2, y, 1)

    J_x = Matrix([[df1_dx, df1_dy], [df2_dx, df2_dy]])

    data = {
        'x': [],
        'y': [],
        'f1(x, y)': [],
        'f2(x, y)': [],
        'Et(%)(x)': [],
        'Et(%)(y)': [],
        'Ea(%)(x)': [],
        'Ea(%)(y)': [],
    }

    iter_num = 0
    true_relative_error_x = abs((x_exact - X[0]) / x_exact)
    true_relative_error_y = abs((y_exact - X[1]) / y_exact)
    approximate_relative_error_x = 10
    approximate_relative_error_y = 10

    errors = {}
    while (
            true_relative_error_x > true_relative_error_cond_x or true_relative_error_y > true_relative_error_cond_y) and (
            approximate_relative_error_x > approximate_relative_error_cond_x or approximate_relative_error_y > approximate_relative_error_cond_y) and iter_num <= iter_num_cond:
        data['x'].append(X[0])
        data['y'].append(X[1])
        data['f1(x, y)'].append(F[0])
        data['f2(x, y)'].append(F[1])
        data['Et(%)(x)'].append(true_relative_error_x)
        data['Et(%)(y)'].append(true_relative_error_y)
        data['Ea(%)(x)'].append(approximate_relative_error_x if iter_num != 0 else '---')
        data['Ea(%)(y)'].append(approximate_relative_error_x if iter_num != 0 else '---')

        df1_dx = N(J_x[0, 0].subs({x: X[0], y: X[1]}))
        df1_dy = N(J_x[0, 1].subs({x: X[0], y: X[1]}))
        df2_dx = N(J_x[1, 0].subs({x: X[0], y: X[1]}))
        df2_dy = N(J_x[1, 1].subs({x: X[0], y: X[1]}))

        J = Matrix([[df1_dx, df1_dy], [df2_dx, df2_dy]])
        J_inv = J.inv()

        Y = J_inv * (-F)

        pre_x_i = X[0]
        pre_y_i = X[1]

        X = X + Y
        F = Matrix([[N(f_1.subs({x: X[0], y: X[1]}))], [N(f_2.subs({x: X[0], y: X[1]}))]])

        iter_num += 1
        true_relative_error_x = abs((x_exact - X[0]) / x_exact)
        true_relative_error_y = abs((y_exact - X[1]) / y_exact)
        approximate_relative_error_x = abs((X[0] - pre_x_i) / X[0])
        approximate_relative_error_y = abs((X[1] - pre_y_i) / X[1])

        errors[iter_num] = [true_relative_error_x, true_relative_error_y, approximate_relative_error_x,
                            approximate_relative_error_y]
    return data, errors


def get_doc(title, f_1, f_2, x_exact, y_exact, x_i, y_i, approximate_relative_error_cond_x, approximate_relative_error_cond_y, true_relative_error_cond_x, true_relative_error_cond_y, iter_num_cond, data):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(f'The solution of this equations:\n"f1(x, y) = {f_1}"\n"f2(x, y) = {f_2}"\ngiving this info\n• x initial: {x_i}      • y initial: {y_i}      • x exact: {x_exact}      • y exact: {y_exact}\n with these conditions\n\t1. true relative error condition: Et(x) < {true_relative_error_cond_x}\n\t2. true relative error condition: Et(y) < {true_relative_error_cond_y}\n\t3. approximate relative error condition: Ea(x) < {approximate_relative_error_cond_x}\n\t4. approximate relative error condition: Ea(y) < {approximate_relative_error_cond_y}\n\t5. iteration number: iter < {iter_num_cond}', style='Intense Quote')

    headers = ['iteration'] + list(data.keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    document.add_picture('chart.png', width=Inches(5))

    document.save('documentation.docx')


st.title('Newton Raphson Method for non linear system')

with st.form("Newton_Raphson_fpr_non_linear_system_form"):
    f_1 = st.text_input("Enter your first function", value=r"x**2 + y - x - 0.4")
    f_1 = sympify(f_1)
    st.latex(f_1)

    f_2 = st.text_input("Enter your second function", value=r"2 * y + 3 * x * y - 2 * x**3")
    f_2 = sympify(f_2)
    st.latex(f_2)

    col1, col2 = st.columns(2)
    with col1:
        x_i = st.text_input("Enter the initial value for x", value=-5)
        x_i = float(x_i)
        st.write('initial x:', x_i)

        x_exact = st.text_input("Enter your exact solution for x", value=-0.55477)
        x_exact = float(x_exact)
        st.write('exact value for x:', x_exact)

        approximate_relative_error_cond_x = st.text_input("Enter the approximate error condition value for x", value=0.005)
        approximate_relative_error_cond_x = float(approximate_relative_error_cond_x)
        st.write('approximate relative error for x:', approximate_relative_error_cond_x)

        true_relative_error_cond_x = st.text_input("Enter the true error condition value for x", value=0.005)
        true_relative_error_cond_x = float(true_relative_error_cond_x)
        st.write('true relative error for x:', true_relative_error_cond_x)

        iter_num_cond = st.text_input("Enter the iteration number condition value", value=10)
        iter_num_cond = int(iter_num_cond)
        st.write('iteration number:', iter_num_cond)
    with col2:
        y_i = st.text_input("Enter the initial value for y", value=0)
        y_i = float(y_i)
        st.write('initial y:', y_i)

        y_exact = st.text_input("Enter your exact solution for y", value=-1.0173242)
        y_exact = float(y_exact)
        st.write('exact value for y:', y_exact)

        approximate_relative_error_cond_y = st.text_input("Enter the approximate error condition value for y",
                                                          value=0.005)
        approximate_relative_error_cond_y = float(approximate_relative_error_cond_y)
        st.write('approximate relative error for y:', approximate_relative_error_cond_y)

        true_relative_error_cond_y = st.text_input("Enter the true error condition value for y", value=0.005)
        true_relative_error_cond_y = float(true_relative_error_cond_y)
        st.write('true relative error for y:', true_relative_error_cond_y)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        data, errors = newton_raphson_for_non_linear_system(f_1, f_2, x_exact, y_exact, x_i, y_i, approximate_relative_error_cond_x, approximate_relative_error_cond_y, true_relative_error_cond_x, true_relative_error_cond_y, iter_num_cond)

        df = pd.DataFrame(data)
        df
        draw_approximate_and_true_errors(errors)

        get_doc('Newton Raphson Method for non linear system', f_1, f_2, x_exact, y_exact, x_i, y_i, approximate_relative_error_cond_x, approximate_relative_error_cond_y, true_relative_error_cond_x, true_relative_error_cond_y, iter_num_cond, data)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")
