import sympy

from sympy import diff, symbols, factorial, sympify, E, N
import streamlit as st
import pandas as pd

from docx import Document
from docx.shared import Inches


def get_doc(title, f_x, x_i, a, num_of_terms, data):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(f'The solution of this equation:\n"{f_x}"\ngiving this info\n• x initial: {x_i}      • a: {a}      • # of terms: {num_of_terms}', style='Intense Quote')

    headers = ['# of terms'] + list(data.keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != '# of terms' else row)

    document.save('documentation.docx')


def taylor(f_x, x_i, a, num_of_terms):
    x, e = symbols("x e")
    f_x = N(f_x.subs(e, E))
    f_x = sympy.sympify(f_x)
    exact_solution = f_x.subs(x, x_i)

    data = {
        'approximate value': [],
        'absolute truncation error': [],
        'relative true error': []
    }

    series = sympy.sympify(0)
    for term in range(0, num_of_terms + 1):
        series += diff(f_x, x, term).subs(x, a) * (x - a) ** term / factorial(term)
        approximate_value = series.subs(x, x_i)

        data['approximate value'].append(approximate_value)
        data['absolute truncation error'].append(abs(exact_solution - approximate_value))
        data['relative true error'].append(abs((exact_solution - approximate_value) / exact_solution))

    return data


st.title('Taylor And Maclaurin Series')

with st.form("Taylor_And_Maclaurin_Series_form"):
    f_x = st.text_input("Enter your function", value=r"e**x")
    f_x = sympify(f_x)
    st.latex(f_x)

    col1, col2 = st.columns(2)
    with col1:
        x_i = st.text_input("Enter the starting point (x)", value=1)
        x_i = float(x_i)
        st.write('x:', x_i)

        num_of_terms = st.text_input("Enter the number of terms", value=5)
        num_of_terms = int(num_of_terms)
        st.write('number of terms:', num_of_terms)

    with col2:
        a = st.text_input("Enter the pivot point (a)", value=0)
        a = float(a)
        st.write('a:', a)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        data = taylor(f_x, x_i, a, num_of_terms)

        df = pd.DataFrame(data)
        df

        get_doc('Taylor And Maclaurin Series', f_x, x_i, a, num_of_terms, data)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")
