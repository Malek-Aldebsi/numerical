import numpy as np
import streamlit as st
import random

from docx.shared import Inches
from sympy import diff, symbols, sympify, lambdify, N
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('TkAgg')
from docx import Document


def draw_fx(f_x, x_test):
    x = symbols("x")
    f_x_test = f_x.subs(x, x_test)

    f_x_for_numpy = lambdify(x, f_x, modules=['numpy'])
    x_values = np.linspace(x_test-3, x_test+3, 100)
    y_values = f_x_for_numpy(x_values)

    plt.plot(x_values, y_values)
    plt.plot(x_test, f_x_test, 'ro', label='Data Point')
    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('x')
    plt.ylabel('f(x)')
    plt.title('My Function')

    plt.savefig('chart.png')
    st.pyplot(plt)


def condition_num(f_x, x_test):
    x = symbols("x")
    f_x_test = f_x.subs(x, x_test)
    f_x_test_d_x = diff(f_x, x, 1).subs(x, x_test)
    x_condition_num = abs((f_x_test_d_x * x_test) / f_x_test)

    return x_condition_num


def get_doc(title, f_x, x_test, x_condition_num):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(f'The condition number of this equation: "{f_x}"\nat this point ({x_test}) = {x_condition_num}', style='Intense Quote')

    document.add_picture('chart.png', width=Inches(5))

    document.save('documentation.docx')


st.title('Condition Number')


with st.form("ConditionNum_form"):
    f_x = st.text_input("Enter your function", value=r"ln(x)")
    f_x = sympify(f_x)
    st.latex(f_x)

    x_test = st.text_input("Enter the value of x", value=1.1)
    x_test = float(x_test)
    st.write('x:', x_test)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        x_condition_num = condition_num(f_x, x_test)
        st.write('result:', x_condition_num)

        draw_fx(f_x, x_test)

        get_doc('Condition Number', f_x, x_test, x_condition_num)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")

# print("f(x)= " + str(coefficientsValues["a7"]) + " * x + " + str(coefficientsValues["a8"]) + " * x^2 + " + str(coefficientsValues["a9"]) + " * x^3")
# print("diff(f(x))= " + str(coefficientsValues["a7"]) + " " + str(coefficientsValues["a8"]) + " * 2 * x " + str(coefficientsValues["a9"]) + " * 3 * x^2")
# print("ConditionNumber = diff(f(x)) * x / f(x)")
#
# print('for x_well:')
# print("f(" + str(x_well) + ") = " + str(f_x_well))
# print("diff(f(" + str(x_well) + "))= " + str(f_x_well_d_x))
# print("Well Condition Number= " + str(x_well_condition_num))
# print('for x_ill')
# print("diff(f(" + str(x_ill) + "))= " + str(f_x_ill_d_x))
# print("f(" + str(x_ill) + ") = " + str(f_x_ill))
# print("Ill Condition Number= " + str(x_ill_condition_num))
