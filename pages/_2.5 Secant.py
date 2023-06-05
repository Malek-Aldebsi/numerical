from sympy import symbols, sympify, N, E

import streamlit as st
import pandas as pd

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('TkAgg')

from docx import Document
from docx.shared import Inches


def draw_approximate_and_true_errors(errors):
    plt.plot(errors.keys(), [error[0] for error in errors.values()], label="Et(%)")
    plt.plot(errors.keys(), [error[1] for error in errors.values()], label="Ea(%)")
    plt.legend()

    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')

    plt.title('True and approximate errors for Secant method')

    plt.savefig('chart.png')
    st.pyplot(plt)


def secant(f_x, x_exact, x_i, x_i_1, approximate_relative_error_cond=0.005, true_relative_error_cond=0.005, iter_num_cond=10):
    x, e = symbols('x e')
    f_x = N(f_x.subs(e, E))

    f_x_i = f_x.subs(x, x_i)
    f_x_i_1 = f_x.subs(x, x_i_1)

    x_i_2 = x_i_1 - (f_x.subs(x, x_i_1) * (x_i - x_i_1)) / (f_x.subs(x, x_i) - f_x.subs(x, x_i_1))

    data = {
        'Xi-1': [],
        'Xi': [],
        'Xi+1': [],
        'f(Xi-1)': [],
        'f(Xi)': [],
        'Et(%)': [],
        'Ea(%)': []
    }

    iter_num = 0
    true_relative_error = abs((x_exact - x_i_2) / x_exact) * 100
    approximate_relative_error = 1
    errors = {}
    while true_relative_error > true_relative_error_cond and approximate_relative_error > approximate_relative_error_cond and iter_num <= iter_num_cond:
        data['Xi-1'].append(x_i)
        data['Xi'].append(x_i_1)
        data['Xi+1'].append(x_i_2)
        data['f(Xi-1)'].append(f_x_i)
        data['f(Xi)'].append(f_x_i_1)
        data['Et(%)'].append(true_relative_error)
        data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

        x_i = x_i_1
        x_i_1 = x_i_2
        x_i_2 = x_i_1 - (f_x.subs(x, x_i_1) * (x_i - x_i_1)) / (f_x.subs(x, x_i) - f_x.subs(x, x_i_1))

        f_x_i = f_x.subs(x, x_i)
        f_x_i_1 = f_x.subs(x, x_i_1)

        iter_num += 1
        true_relative_error = abs((x_exact - x_i_2) / x_exact) * 100
        approximate_relative_error = abs((x_i_2 - x_i_1) / x_i_2) * 100

        errors[iter_num] = [true_relative_error, approximate_relative_error]

    return data, errors


def get_doc(title, f_x, x_exact, x_i, x_i_1, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond, data):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(f'The solution of this equation:\n"{f_x}"\ngiving this info\n• Xi initial: {x_i_1}      • Xi-1: {x_i}      • x exact: {x_exact}\n with these conditions\n\t1. true relative error condition: Et < {true_relative_error_cond}\n\t2. approximate relative error condition: Ea < {approximate_relative_error_cond}\n\t3. iteration number: iter < {iter_num_cond}', style='Intense Quote')

    headers = ['iteration'] + list(data.keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    document.add_picture('chart.png', width=Inches(5))

    document.save('documentation.docx')


st.title('Secant Method')

with st.form("Secant_form"):
    f_x = st.text_input("Enter your function", value=r"e**(-x) - x")
    f_x = sympify(f_x)
    st.latex(f_x)

    col1, col2 = st.columns(2)
    with col1:
        x_i = st.text_input("Enter the first initial value", value=0)
        x_i = float(x_i)
        st.write('initial Xi-1:', x_i)

        x_exact = st.text_input("Enter your exact solution", value=0.56714329)
        x_exact = float(x_exact)
        st.write('exact value:', x_exact)

        approximate_relative_error_cond = st.text_input("Enter the approximate error condition value", value=0.005)
        approximate_relative_error_cond = float(approximate_relative_error_cond)
        st.write('approximate relative error:', approximate_relative_error_cond)

    with col2:
        x_i_1 = st.text_input("Enter the second initial value", value=1)
        x_i_1 = float(x_i_1)
        st.write('initial Xi:', x_i_1)

        iter_num_cond = st.text_input("Enter the iteration number condition value", value=3)
        iter_num_cond = int(iter_num_cond)
        st.write('iteration number:', iter_num_cond)

        true_relative_error_cond = st.text_input("Enter the true error condition value", value=0.005)
        true_relative_error_cond = float(true_relative_error_cond)
        st.write('true relative error:', true_relative_error_cond)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        data, errors = secant(f_x, x_exact, x_i, x_i_1, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond)

        df = pd.DataFrame(data)
        df
        draw_approximate_and_true_errors(errors)

        get_doc('Secant Method', f_x, x_exact, x_i, x_i_1, approximate_relative_error_cond,
                true_relative_error_cond, iter_num_cond, data)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")
