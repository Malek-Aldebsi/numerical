from sympy import symbols, sympify, N, E

import streamlit as st
import pandas as pd

import matplotlib.pyplot as plt
import matplotlib

from docx import Document
from docx.shared import Inches


def draw_approximate_and_true_errors(errors):
    plt.plot(errors.keys(), [error[0] for error in errors.values()], label="Et(%)")
    plt.plot(errors.keys(), [error[1] for error in errors.values()], label="Ea(%)")
    plt.legend()

    plt.grid(True, linestyle='--', color='gray', linewidth=0.5)
    plt.xlabel('iteration')
    plt.ylabel('error(%)')

    plt.title('True and approximate errors for False Position method')

    plt.savefig('chart.png')
    st.pyplot(plt)


def false_position(f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond=0.005, true_relative_error_cond=0.005, iter_num_cond=10):
    x, e = symbols('x e')
    f_x = N(f_x.subs(e, E))

    f_x_l = f_x.subs(x, x_lower)
    f_x_u = f_x.subs(x, x_upper)

    if f_x_u * f_x_l > 0:
        st.error('lower and upper not bracketing the exact solution')

    else:
        x_root = x_upper - f_x_u * (x_upper - x_lower) / (f_x_u - f_x_l)
        f_x_r = f_x.subs(x, x_root)

        data = {
            'x lower': [],
            'x upper': [],
            'x root': [],
            'f(xr)': [],
            'Et(%)': [],
            'Ea(%)': []
        }

        iter_num = 0
        true_relative_error = abs((x_exact - x_root) / x_exact) * 100
        approximate_relative_error = 1
        errors = {}
        while true_relative_error > true_relative_error_cond and approximate_relative_error > approximate_relative_error_cond and iter_num <= iter_num_cond:
            data['x lower'].append(x_lower)
            data['x upper'].append(x_upper)
            data['x root'].append(x_root)
            data['f(xr)'].append(f_x_r)
            data['Et(%)'].append(true_relative_error)
            data['Ea(%)'].append(approximate_relative_error if iter_num != 0 else '---')

            if f_x_r * f_x_u < 0:
                x_lower = x_root
                f_x_l = f_x_r
            else:
                x_upper = x_root
                f_x_u = f_x_r

            pre_x_r = x_root
            x_root = x_upper - f_x_u * (x_upper - x_lower) / (f_x_u - f_x_l)
            f_x_r = f_x.subs(x, x_root)

            iter_num += 1
            true_relative_error = abs((x_exact - x_root) / x_exact) * 100
            approximate_relative_error = abs((x_root - pre_x_r) / x_root) * 100

            errors[iter_num] = [true_relative_error, approximate_relative_error]

        return data, errors


def get_doc(title, f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond, data):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(f'The solution of this equation:\n"{f_x}"\ngiving this info\n• x lower: {x_lower}      • x upper: {x_upper}      • x exact: {x_exact}\n with these conditions\n\t1. true relative error condition: Et < {true_relative_error_cond}\n\t2. approximate relative error condition: Ea < {approximate_relative_error_cond}\n\t3. iteration number: iter < {iter_num_cond}', style='Intense Quote')

    headers = ['iteration'] + list(data.keys())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in range(len(list(data.values())[0])):
        row_cells = table.add_row().cells
        for index, column in enumerate(headers):
            row_cells[index].text = str(data[column][row] if column != 'iteration' else row)

    document.add_picture('chart.png', width=Inches(5))

    document.save('documentation.docx')


st.title('False Position Method')

with st.form("False_Position_form"):
    f_x = st.text_input("Enter your function", value=r"4.84 * x**3 + 3.59 * x**2 + 2.86 * x +0.134")
    f_x = sympify(f_x)
    st.latex(f_x)

    col1, col2 = st.columns(2)
    with col1:
        x_lower = st.text_input("Enter the lower bracket", value=-0.5)
        x_lower = float(x_lower)
        st.write('Lower:', x_lower)

        x_exact = st.text_input("Enter your exact solution", value=-0.05)
        x_exact = float(x_exact)
        st.write('exact value:', x_exact)

        approximate_relative_error_cond = st.text_input("Enter the approximate error condition value", value=0.005)
        approximate_relative_error_cond = float(approximate_relative_error_cond)
        st.write('approximate relative error:', approximate_relative_error_cond)

    with col2:
        x_upper = st.text_input("Enter the upper bracket", value=0.5)
        x_upper = float(x_upper)
        st.write('Upper:', x_upper)

        iter_num_cond = st.text_input("Enter the iteration number condition value", value=10)
        iter_num_cond = int(iter_num_cond)
        st.write('iteration number:', iter_num_cond)

        true_relative_error_cond = st.text_input("Enter the true error condition value", value=0.005)
        true_relative_error_cond = float(true_relative_error_cond)
        st.write('true relative error:', true_relative_error_cond)

    submit_button = st.form_submit_button(label="Submit")

    if submit_button:
        data, errors = false_position(f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond, true_relative_error_cond, iter_num_cond)

        df = pd.DataFrame(data)
        df
        draw_approximate_and_true_errors(errors)

        get_doc('False Position Method', f_x, x_exact, x_lower, x_upper, approximate_relative_error_cond,
                true_relative_error_cond, iter_num_cond, data)

st.download_button("Download Document", data=open('documentation.docx', "rb"), file_name="document.docx")
